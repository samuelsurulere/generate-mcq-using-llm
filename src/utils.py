from langchain.callbacks.manager import get_openai_callback
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import json



# Define prompt templates
first_template = """
VOCAB:
{vocab_handout}
You are an examiner with the Apprenticeship and Industry Training (AIT) organization in Alberta, Canada. The {vocab_handout} is a summmary of science concepts 
that is necessary to prepare preapprenticeship candidates for the entrance exam. You have been tasked with generating multiple choice questions for a 
preapprenticeship exam in the science subject. Using the {vocab_handout}, please generate mutiple choice questions having {number} questions. The quiz questions 
should conform to the standard of the AIT exam questions. Make sure the questions are not repeated and do format your response using the {response_json} template 
as a guide. Please include some questions that will involve calculations and also conversion between units. The tone of the multiple choice questions should be similar to 
the preapprenticeship exam sample questions. Most of the questions should require the candidates to think critically. Ensure that you cross-check the {vocab_handout}
and provide the correct information for each question and the exact correct answer for each question. Use the {response_json} template below to format your response. 
The total number of quesions must be {number}.
"""

second_template = """
You are an expert AIT exam reviewer and you ensure that exam questions meet up to the required standard for preapprenticeship entrance exams. Given some
multiple choice questions recently generated, you need to evaluate the complexity of the question and give a complete analysis of the {quiz}. Only use a
maximum of 100 words for the complexity analysis. If the quiz is not at par with the required standard, update the {quiz} question that needs to be changed.
Check for the correctness of the questions, formulas, units and also for the correct answer.
Quiz_MCQs:
{quiz}

Check from an expert reviewer for the above quiz:
"""

# Create a new Word document
doc = Document()
doc.add_heading('Sample Practice Quiz Questions', level=1)


def generate_quiz_from_pdf(pdf_file, vocab_handout, number, response_json, generate_evaluate_chain):
    """
    Generates quiz questions from a PDF file and appends them to a Word document.

    Args:
        pdf_file (str): Path to the PDF file.
        vocab_handout (str): Vocabulary handout text.
        number (int): Number of questions to generate.
        response_json (str): JSON-formatted string for quiz response template.
    """
    # Read PDF file
    read_vocab = PdfReader(pdf_file)
    for i, page in enumerate(read_vocab.pages):
        content = page.extract_text()
        if content:
            vocab_handout += content
    
    # Generate quiz questions
    with get_openai_callback() as cb:
        response = generate_evaluate_chain({
            "vocab_handout": vocab_handout,
            "number": number,
            "response_json": response_json,
        })
    
    print(f"Total Tokens: {cb.total_tokens}")
    print(f"Prompt Tokens: {cb.prompt_tokens}")
    print(f"Completion Tokens: {cb.completion_tokens}")
    print(f"Total Cost: {cb.total_cost}")
    
    # Extract quiz questions
    quiz = response.get("quiz")
    quiz = json.loads(quiz)
    
    # Append quiz questions to Word document
    for _, value in quiz.items():
        mcq = value["mcq"]
        options = "\n".join([f"{option}: {option_value}" for option, option_value in value["options"].items()])
        correct = value["correct"]
        
        # Add question to Word document
        doc.add_paragraph(f'{mcq}', style='BodyText')
        doc.add_paragraph(options, style='BodyText')
        doc.add_paragraph(f'Correct Answer: {correct}', style='BodyText')
        
        # Add spacing between questions
        doc.add_paragraph('\n')


# Define JSON response template for quiz questions
RESPONSE_JSON = {
    "1": {
        "mcq": "multiple choice question",
        "options": {
            "A.": "choice here",
            "B.": "choice here",
            "C.": "choice here",
            "D.": "choice here",
        },
        "correct": "correct answer",
    },
    "2": {
        "mcq": "multiple choice question",
        "options": {
            "A.": "choice here",
            "B.": "choice here",
            "C.": "choice here",
            "D.": "choice here",
        },
        "correct": "correct answer",
    },
    "3": {
        "mcq": "multiple choice question",
        "options": {
            "A.": "choice here",
            "B.": "choice here",
            "C.": "choice here",
            "D.": "choice here",
        },
        "correct": "correct answer",
    },
}

